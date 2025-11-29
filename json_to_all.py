import sys
import json
import uuid
import os
from pathlib import Path
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from openpyxl import Workbook, load_workbook

# ---------------------------------
# 1. Read JSON from file path argument
# ---------------------------------
if len(sys.argv) < 2:
    print(json.dumps({"error": "No JSON file path provided"}))
    sys.exit(1)

json_file_path = sys.argv[1]

try:
    with open(json_file_path, "r", encoding="utf-8") as f:
        data = json.load(f)
except Exception as e:
    print(json.dumps({"error": f"Failed to read JSON file: {e}"}))
    sys.exit(1)

# ---------------------------------
# 2. Prepare paths and unique ID
# ---------------------------------
BASE_DIR = Path(__file__).resolve().parent
output_dir = BASE_DIR / "uploads" / "reports"
output_dir.mkdir(parents=True, exist_ok=True)

unique_id = str(uuid.uuid4())[:8]  # e.g. "a3f9c12b"

pdf_path = output_dir / f"{unique_id}.pdf"
docx_path = output_dir / f"{unique_id}.docx"
excel_path = output_dir / "database.xlsx"

# Expected keys (in fixed order for Excel)
expected_keys = [
    "parties",
    "incident_summary",
    "key_facts",
    "evidence",
    "ipc_sections",
    "timeline",
    "red_flags",
    "recommendations",
    "final_summary",
]

# Normalize: make sure all keys exist
for k in expected_keys:
    if k not in data:
        if k in ("incident_summary", "final_summary"):
            data[k] = ""
        else:
            data[k] = []

# ---------------------------------
# 3. Generate DOCX
# ---------------------------------
doc = Document()
doc.add_heading("Legal Document Analysis", level=1)

for key in expected_keys:
    value = data.get(key, "")

    doc.add_heading(key.replace("_", " ").title(), level=2)

    if isinstance(value, list):
        for item in value:
            doc.add_paragraph(f"- {item}", style="List Bullet")
    else:
        doc.add_paragraph(str(value))

doc.save(str(docx_path))

# ---------------------------------
# 4. Generate PDF
# ---------------------------------
c = canvas.Canvas(str(pdf_path), pagesize=letter)
width, height = letter
y = height - 50

c.setFont("Helvetica-Bold", 16)
c.drawString(50, y, "Legal Document Analysis")
y -= 30

def write_line(text, y_pos):
    """
    Draws one line and returns updated y.
    """
    text = text.strip()
    if not text:
        return y_pos

    c.drawString(70, y_pos, text)
    y_pos -= 15
    if y_pos < 50:
        c.showPage()
        y_pos = height - 50
        c.setFont("Helvetica", 11)
    return y_pos

for key in expected_keys:
    value = data.get(key, "")

    if y < 80:
        c.showPage()
        y = height - 50

    c.setFont("Helvetica-Bold", 13)
    c.drawString(50, y, key.replace("_", " ").title())
    y -= 20

    c.setFont("Helvetica", 11)

    if isinstance(value, list):
        for item in value:
            y = write_line(f"- {item}", y)
    else:
        text_str = str(value)
        for line in text_str.split(". "):
            y = write_line(line, y)

c.save()

# ---------------------------------
# 5. Generate / append Excel (DB style)
# ---------------------------------
columns = expected_keys

if not excel_path.exists():
    wb = Workbook()
    ws = wb.active
    ws.title = "Cases"
    ws.append(columns)
else:
    wb = load_workbook(str(excel_path))
    ws = wb.active

row = []
for col in columns:
    val = data.get(col, "")
    if isinstance(val, list):
        val = "; ".join(str(v) for v in val)
    row.append(val)

ws.append(row)
wb.save(str(excel_path))

# ---------------------------------
# 6. Print JSON result
# ---------------------------------
print(json.dumps({
    "pdf_path": str(pdf_path),
    "docx_path": str(docx_path),
    "excel_path": str(excel_path),
    "id": unique_id
}))
